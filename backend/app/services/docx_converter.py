# NOTICE: This file is protected under RCF-PL v2.0.3
"""WRT format converter — bidirectional .docx ↔ .wrt.

.wrt is a lightweight tagged-text format designed for AI editing of
office documents.  The agent works with .wrt (plain text + semantic
tags), while the user downloads/uploads real .docx files.

Format specification (v2):
    [h1]...[/h1]  [h2]...[/h2]  [h3]...[/h3]  — headings
    [b]...[/b]    — bold
    [i]...[/i]    — italic
    [u]...[/u]    — underline
    [s]...[/s]    — strikethrough
    [code]...[/code] — monospace / code
    [quote]...[/quote] — block quote
    [list]         — unordered list (items start with * )
    [/list]        — end list
    [table]        — table (pipe-delimited rows)
    [/table]       — end table
    [img src="data:image/png;base64,..." alt="..."] — embedded image (base64)
    [img alt="..."] — image placeholder (for compatibility)

    Everything outside tags is plain paragraph text.
    Blank lines separate paragraphs.
"""
from __future__ import annotations

import base64
import io
import logging
import re
from typing import Any

log = logging.getLogger(__name__)


# ── .docx → .wrt ────────────────────────────────────────────────────────────

def _para_tag(para: Any) -> str | None:
    """Map a python-docx paragraph style to a .wrt tag."""
    style = (para.style.name or "").lower()
    if "heading 1" in style or style == "title":
        return "h1"
    if "heading 2" in style:
        return "h2"
    if "heading 3" in style or "heading" in style:
        return "h3"
    if "quote" in style or "block" in style:
        return "quote"
    if "list" in style:
        return None  # handled via list items
    return None


def _run_markup(run: Any) -> tuple[str, list[str]]:
    """Wrap a single run's text with .wrt inline tags.

    Returns:
        (text, images): text with inline tags, and list of base64 image lines
    """
    text = run.text or ""
    images: list[str] = []

    # Check for inline images (drawing elements)
    if hasattr(run, "_element") and run._element is not None:
        for drawing in run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"):
            # Extract image data from drawing
            img_line = _extract_image_from_drawing(drawing, run)
            if img_line:
                images.append(img_line)

    if not text and not images:
        return "", []

    if run.bold:
        text = f"[b]{text}[/b]"
    if run.italic:
        text = f"[i]{text}[/i]"
    if run.underline:
        text = f"[u]{text}[/u]"
    if run.font.strike:
        text = f"[s]{text}[/s]"
    # Monospace / code font
    font_name = (run.font.name or "").lower()
    if "mono" in font_name or "courier" in font_name or "code" in font_name:
        if not run.bold and not run.italic:
            text = f"[code]{text}[/code]"
    return text, images


def _extract_image_from_drawing(drawing: Any, run: Any) -> str | None:
    """Extract image from a drawing element and encode as base64.

    Returns:
        .wrt image tag with base64 data, or None if extraction fails
    """
    try:
        # Find blip element (contains image reference)
        blip = drawing.find(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
        if blip is None:
            return None

        # Get image relationship ID
        embed_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if not embed_id:
            return None

        # Get image part from document
        doc_part = run.part
        image_part = doc_part.related_parts[embed_id]
        image_bytes = image_part.blob

        # Detect image type from content_type
        content_type = image_part.content_type
        if "png" in content_type:
            mime = "image/png"
        elif "jpeg" in content_type or "jpg" in content_type:
            mime = "image/jpeg"
        elif "gif" in content_type:
            mime = "image/gif"
        elif "bmp" in content_type:
            mime = "image/bmp"
        else:
            mime = "image/png"  # fallback

        # Encode to base64
        b64 = base64.b64encode(image_bytes).decode("utf-8")

        # Extract alt text from drawing description if available
        alt_text = "image"
        desc_elem = drawing.find(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr")
        if desc_elem is not None:
            alt_text = desc_elem.get("descr", "image") or desc_elem.get("title", "image") or "image"

        return f'[img src="data:{mime};base64,{b64}" alt="{alt_text}"]'

    except Exception as e:
        log.warning(f"Failed to extract image from drawing: {e}")
        return None


def docx_to_wrt(docx_bytes: bytes) -> str:
    """Convert a .docx file (bytes) to .wrt text format with embedded images."""
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    lines: list[str] = []

    for para in doc.paragraphs:
        # Check if this is a list item
        pstyle = (para.style.name or "").lower()
        is_list = "list" in pstyle

        if is_list:
            # Build inline markup from runs, collecting images
            text_parts: list[str] = []
            for run in para.runs:
                text, images = _run_markup(run)
                if text:
                    text_parts.append(text)
                for img_line in images:
                    lines.append(img_line)
            inner = "".join(text_parts)
            if inner.strip():
                lines.append(f"* {inner.strip()}")
            continue

        # Regular paragraph or heading
        tag = _para_tag(para)
        text_parts: list[str] = []
        for run in para.runs:
            text, images = _run_markup(run)
            if text:
                text_parts.append(text)
            # Images go on separate lines
            for img_line in images:
                lines.append(img_line)

        inner = "".join(text_parts)
        if not inner.strip():
            # Paragraph might have only images
            if not any(img for _, imgs in [_run_markup(r) for r in para.runs] for img in imgs):
                lines.append("")
            continue

        if tag:
            lines.append(f"[{tag}]{inner.strip()}[/{tag}]")
        else:
            lines.append(inner.strip())

    # Extract tables
    for table in doc.tables:
        lines.append("[table]")
        for row in table.rows:
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("[/table]")

    # Clean up multiple blank lines
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip() + "\n"


# ── .wrt → .docx ────────────────────────────────────────────────────────────

_TAG_RE = re.compile(
    r"\[(b|i|u|s|h1|h2|h3|code|quote)\](.*?)\[/\1\]",
    re.DOTALL,
)

# Regex for image tags with base64 data
_IMG_TAG_RE = re.compile(
    r'\[img\s+src="data:image/(png|jpeg|jpg|gif|bmp);base64,([^"]+)"\s+alt="([^"]*)"\]',
    re.IGNORECASE,
)

# Regex for legacy image placeholder tags (no base64)
_IMG_PLACEHOLDER_RE = re.compile(
    r'\[img\s+alt="([^"]*)"\]',
    re.IGNORECASE,
)


def _apply_inline_tags(doc: Any, paragraph: Any, text: str) -> None:
    """Parse inline .wrt tags and add runs to a paragraph."""
    pos = 0
    for m in _TAG_RE.finditer(text):
        # Plain text before the tag
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])

        tag = m.group(1)
        inner = m.group(2)
        run = paragraph.add_run(inner)

        if tag == "b":
            run.bold = True
        elif tag == "i":
            run.italic = True
        elif tag == "u":
            run.underline = True
        elif tag == "s":
            run.font.strike = True
        elif tag == "code":
            run.font.name = "Courier New"
            run.font.size = __import__("docx.shared").Pt(9)
        elif tag in ("h1", "h2", "h3"):
            run.bold = True

        pos = m.end()

    # Remaining plain text
    if pos < len(text):
        paragraph.add_run(text[pos:])


def wrt_to_docx(wrt_text: str) -> bytes:
    """Convert .wrt text format to a .docx file (bytes) with embedded images."""
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    lines = wrt_text.split("\n")

    i = 0
    in_table = False
    in_list = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Image tag with base64 data
        img_match = _IMG_TAG_RE.match(stripped)
        if img_match:
            img_format = img_match.group(1)  # png, jpeg, jpg, gif, bmp
            b64_data = img_match.group(2)
            alt_text = img_match.group(3)

            try:
                # Decode base64
                img_bytes = base64.b64decode(b64_data)
                img_stream = io.BytesIO(img_bytes)

                # Add image to document (max width 6 inches to fit page)
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(img_stream, width=Inches(6))
            except Exception as e:
                log.warning(f"Failed to insert image: {e}")
                # Fallback: insert as placeholder text
                doc.add_paragraph(f"[Image: {alt_text}]")

            i += 1
            continue

        # Legacy image placeholder (no base64)
        placeholder_match = _IMG_PLACEHOLDER_RE.match(stripped)
        if placeholder_match:
            alt_text = placeholder_match.group(1)
            doc.add_paragraph(f"[Image: {alt_text}]")
            i += 1
            continue

        # Table
        if stripped == "[table]":
            in_table = True
            i += 1
            continue
        if stripped == "[/table]":
            in_table = False
            i += 1
            continue
        if in_table and stripped.startswith("|"):
            # Parse table row — first row = header
            cells = [c.strip().replace("\\|", "|") for c in stripped.split("|")[1:-1]]
            if not doc.tables:
                doc.add_table(rows=1, cols=max(len(cells), 1))
            table = doc.tables[-1]
            if len(table.rows) == 1 and all(c == "" for c in cells):
                i += 1
                continue
            row = table.add_row()
            for j, cell_text in enumerate(cells):
                row.cells[j].text = cell_text
            i += 1
            continue

        # List
        if stripped == "[list]":
            in_list = True
            i += 1
            continue
        if stripped == "[/list]":
            in_list = False
            i += 1
            continue
        if in_list and stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            i += 1
            continue

        # Headings
        heading_match = re.match(r"\[(h[1-3])\](.*?)\[/\1\]", stripped)
        if heading_match:
            tag = heading_match.group(1)
            inner = heading_match.group(2)
            level = int(tag[1])
            doc.add_heading(inner, level=level)
            i += 1
            continue

        # Block quote
        if stripped.startswith("[quote]") and stripped.endswith("[/quote]"):
            inner = stripped[7:-8]
            doc.add_paragraph(inner, style="Quote")
            i += 1
            continue

        # Empty line = paragraph break
        if not stripped:
            i += 1
            continue

        # Regular paragraph with inline tags
        para = doc.add_paragraph()
        _apply_inline_tags(doc, para, stripped)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
