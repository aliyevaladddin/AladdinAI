# NOTICE: This file is protected under RCF-PL v2.0.3
"""Tests for .docx ↔ .wrt image round-trip support."""
import base64
import io

import pytest
from docx import Document
from PIL import Image

from app.services.docx_converter import docx_to_wrt, wrt_to_docx


def create_test_image(width: int = 100, height: int = 100, color: str = "red") -> bytes:
    """Create a simple test image as PNG bytes."""
    img = Image.new("RGB", (width, height), color)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def test_docx_to_wrt_extracts_images():
    """Test that docx_to_wrt extracts images as base64."""
    # Create a .docx with text and an image
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("Paragraph before image.")

    # Add image to document
    img_bytes = create_test_image(50, 50, "blue")
    img_stream = io.BytesIO(img_bytes)
    doc.add_picture(img_stream, width=pytest.importorskip("docx.shared").Inches(2))

    doc.add_paragraph("Paragraph after image.")

    # Convert to bytes
    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    docx_bytes = docx_buf.getvalue()

    # Convert to .wrt
    wrt = docx_to_wrt(docx_bytes)

    # Verify structure
    assert "[h1]Test Document[/h1]" in wrt
    assert "Paragraph before image." in wrt
    assert "Paragraph after image." in wrt

    # Verify image tag with base64 data
    assert '[img src="data:image/' in wrt
    assert ";base64," in wrt
    assert 'alt=' in wrt


def test_wrt_to_docx_embeds_images():
    """Test that wrt_to_docx embeds base64 images back into .docx."""
    # Create test image
    img_bytes = create_test_image(60, 60, "green")
    b64_data = base64.b64encode(img_bytes).decode("utf-8")

    # Create .wrt with embedded image
    wrt_text = f"""[h1]Test Document[/h1]

Text before image.

[img src="data:image/png;base64,{b64_data}" alt="test image"]

Text after image.
"""

    # Convert to .docx
    docx_bytes = wrt_to_docx(wrt_text)

    # Parse resulting .docx
    doc = Document(io.BytesIO(docx_bytes))

    # Verify structure
    assert len(doc.paragraphs) >= 3
    assert doc.paragraphs[0].text == "Test Document"

    # Find paragraphs with text
    text_paras = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "Text before image." in text_paras
    assert "Text after image." in text_paras

    # Verify image was embedded
    has_image = False
    for para in doc.paragraphs:
        if para._element.xpath('.//w:drawing'):
            has_image = True
            break

    assert has_image


def test_image_round_trip():
    """Test full round-trip: .docx → .wrt → .docx preserves images."""
    # Create original .docx with image
    doc1 = Document()
    doc1.add_heading("Round-trip Test", level=1)
    doc1.add_paragraph("Before image")

    img_bytes = create_test_image(40, 40, "yellow")
    img_stream = io.BytesIO(img_bytes)
    doc1.add_picture(img_stream, width=pytest.importorskip("docx.shared").Inches(1.5))

    doc1.add_paragraph("After image")

    # Convert to bytes
    docx1_buf = io.BytesIO()
    doc1.save(docx1_buf)
    docx1_bytes = docx1_buf.getvalue()

    # Round-trip: .docx → .wrt → .docx
    wrt = docx_to_wrt(docx1_bytes)
    docx2_bytes = wrt_to_docx(wrt)

    # Parse result
    doc2 = Document(io.BytesIO(docx2_bytes))

    # Verify text content preserved
    assert doc2.paragraphs[0].text == "Round-trip Test"
    text_paras = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert "Before image" in text_paras
    assert "After image" in text_paras

    # Verify .wrt has base64 image
    assert '[img src="data:image/' in wrt
    assert ";base64," in wrt


def test_wrt_legacy_placeholder():
    """Test that legacy [img alt="..."] placeholders still work."""
    wrt_text = """[h1]Legacy Format[/h1]

Text before.

[img alt="legacy placeholder"]

Text after.
"""

    # Convert to .docx
    docx_bytes = wrt_to_docx(wrt_text)
    doc = Document(io.BytesIO(docx_bytes))

    # Verify placeholder text appears
    text_paras = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("[Image: legacy placeholder]" in p for p in text_paras)


def test_multiple_images():
    """Test document with multiple images."""
    img1 = create_test_image(30, 30, "red")
    img2 = create_test_image(40, 40, "blue")

    b64_1 = base64.b64encode(img1).decode("utf-8")
    b64_2 = base64.b64encode(img2).decode("utf-8")

    wrt_text = f"""[h1]Multiple Images[/h1]

First image:
[img src="data:image/png;base64,{b64_1}" alt="red square"]

Second image:
[img src="data:image/png;base64,{b64_2}" alt="blue square"]

End.
"""

    # Convert to .docx
    docx_bytes = wrt_to_docx(wrt_text)
    doc = Document(io.BytesIO(docx_bytes))

    # Verify structure
    assert doc.paragraphs[0].text == "Multiple Images"
    text_paras = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "First image:" in text_paras
    assert "Second image:" in text_paras
    assert "End." in text_paras


def test_invalid_base64_graceful_fallback():
    """Test that invalid base64 data falls back to placeholder."""
    wrt_text = """[h1]Invalid Image[/h1]

[img src="data:image/png;base64,INVALID_BASE64_DATA!!!" alt="broken"]

Done.
"""

    # Should not raise, should insert placeholder
    docx_bytes = wrt_to_docx(wrt_text)
    doc = Document(io.BytesIO(docx_bytes))

    # Verify placeholder appears
    text_paras = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("[Image: broken]" in p for p in text_paras)
